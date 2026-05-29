"""
Generates asbl-qa-agent-version3.docx
Comprehensive implementation document — May 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

DARK_BLUE  = RGBColor(0x1A, 0x37, 0x6C)
MID_BLUE   = RGBColor(0x1F, 0x5C, 0x99)
ACCENT     = RGBColor(0xE8, 0x6C, 0x1E)
GREEN      = RGBColor(0x1A, 0x7A, 0x3C)
RED        = RGBColor(0xC0, 0x20, 0x20)
LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
TABLE_HDR  = RGBColor(0x1A, 0x37, 0x6C)
GREEN_HDR  = RGBColor(0x1A, 0x7A, 0x3C)
RED_HDR    = RGBColor(0xC0, 0x20, 0x20)
ORANGE_HDR = RGBColor(0xE8, 0x69, 0x1E)

def rgb_to_hex(rgb):
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_to_hex(rgb))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), kwargs.get('val', 'single'))
        border.set(qn('w:sz'), kwargs.get('sz', '4'))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), kwargs.get('color', 'FFFFFF'))
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = DARK_BLUE
        run.font.bold = True
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '1A376C')
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = MID_BLUE
        run.font.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT
        run.font.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_body(doc, text, bold=False, italic=False, indent=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color if color else BLACK
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    return p

def add_bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = color if color else BLACK
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    return p

def add_table(doc, headers, rows, col_widths=None, hdr_color=None):
    hdr_rgb = hdr_color if hdr_color else TABLE_HDR
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, hdr_rgb)
        set_cell_border(cell, val='single', sz='4', color='FFFFFF')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = LIGHT_GREY if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_border(cell, val='single', sz='2', color='CCCCCC')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table

def add_notice(doc, text, color_hex='1A376C'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)
    run.font.bold = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    fill_map = {'1A376C': 'D6E4F7', '1A7A3C': 'E8F5E9', 'C02020': 'FDECEA', 'E8691E': 'FFF3E0'}
    shd.set(qn('w:fill'), fill_map.get(color_hex, 'F5F5F5'))
    pPr.append(shd)


# ════════════════════════════════════════════════════════════════
#  COVER
# ════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ASBL QA Agent System')
run.font.size = Pt(30); run.font.bold = True; run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complete Implementation Document')
run.font.size = Pt(17); run.font.color.rgb = MID_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Everything that is built · How every part works · What is left and why')
run.font.size = Pt(11); run.font.italic = True; run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run = p.add_run('May 2026  |  Confidential — ASBL Internal')
run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
add_heading(doc, 'Table of Contents', 1)
toc = [
    ('1',   'Why This System Exists'),
    ('2',   'How the System Works — Big Picture'),
    ('3',   'The Technology Stack'),
    ('4',   'Agent 1 — Chatbot QA'),
    ('5',   'Agent 2 — Voice QA (Anandita)'),
    ('6',   'Agent 3 — Analytics Validator'),
    ('7',   'Agent 4 — Feedback Agent'),
    ('8',   'Agent 5 — Recommendation Agent'),
    ('9',   'The Runner — How Everything Connects'),
    ('10',  'The Prompt Design Philosophy'),
    ('11',  'The Knowledge Base'),
    ('12',  'Production Database Reference'),
    ('13',  'Results Database — Full Schema'),
    ('14',  'The LLM Wrapper'),
    ('15',  'Notifications — Teams Webhook'),
    ('16',  'Business Outcomes Configuration'),
    ('17',  'ASBL Project Facts — Ground Truth'),
    ('18',  'File & Folder Structure'),
    ('19',  'How to Run'),
    ('20',  'What Is Built vs What Is Pending'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'  {num}.  {title}')
    run.font.size = Pt(10); run.font.color.rgb = BLACK

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 1 — WHY THIS SYSTEM EXISTS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '1.  Why This System Exists', 1)

add_body(doc,
    'ASBL operates two AI products that speak to potential property buyers every day: '
    'a website chatbot and a voice agent called Anandita. Both are always on — they answer '
    'questions about projects, prices, RERA numbers, handover dates, and rental offers at '
    'any hour, without a human on the line.')

add_body(doc,
    'The risk that comes with always-on AI is straightforward: the bot can give a wrong price '
    'to a serious buyer, quote a RERA number incorrectly, describe a feature that does not '
    'exist, or miss a lead entirely. Any of these can cost ASBL a sale or create a legal '
    'liability. And with hundreds of conversations per day, no human can review them all.')

add_body(doc, 'This QA system exists to solve that problem automatically.', bold=True)

add_body(doc,
    'Every 4 hours, without any human effort, the system reads every new chatbot conversation '
    'and every new Anandita call, checks each one for errors against the knowledge base, '
    'validates whether the analytics scoring system is correctly identifying high-intent leads, '
    'computes health scores across all three products, and generates a specific, actionable fix '
    'for every problem it finds — complete with the exact file to open and what to change.')

add_notice(doc,
    'The system catches pricing errors before they compound, flags wasted Meta ad spend '
    'in real time, and ensures Anandita speaks the right language on every call — '
    'all running on a local Mac with a free local LLM. Zero cloud API cost.',
    'E8691E')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 2 — BIG PICTURE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '2.  How the System Works — Big Picture', 1)

add_body(doc,
    'The system is a pipeline of five specialised agents that run in sequence every 4 hours. '
    'Each agent has exactly one job. No agent does more than that job. '
    'The output of each agent feeds into the next.')

add_table(doc,
    ['Step', 'Agent', 'One-Line Job', 'Input', 'Output'],
    [
        ['1', 'Chatbot QA',     'Evaluate every new chatbot conversation for accuracy and quality',     'Production chatbot conversations + knowledge base', 'Pass/Fail per conversation with issues and score'],
        ['2', 'Voice QA',       'Evaluate every new Anandita call for accuracy and quality',            'Production call transcripts + knowledge base',       'Pass/Fail per call with issues, score, language compliance'],
        ['3', 'Analytics',      'Check whether the scoring system is producing the right business outcomes', 'Production analytics scores and Meta events',   'Flags on band rates, multiplier effectiveness, Meta accuracy'],
        ['4', 'Feedback',       'Compute health scores and identify every real problem across all agents', 'All QA results from steps 1-3 + health history',  'Health scores + list of problems with evidence'],
        ['5', 'Recommendation', 'Generate one specific, actionable fix per problem',                    'Problem list from step 4',                           'Ranked list of fixes with exact file paths and expected outcomes'],
    ],
    col_widths=[0.4, 1.4, 2.2, 1.6, 1.6],
)

add_heading(doc, 'Two MongoDB Clusters — Production Is Never Written To', 2)
add_body(doc,
    'The system is built on a strict separation: one MongoDB cluster for reading production '
    'data, a completely separate cluster for writing all QA results. The production cluster '
    'is configured as read-only at the connection level. There is no scenario in which a QA '
    'agent can write to or modify any production data.')

add_table(doc,
    ['Cluster', 'Environment Variable', 'Direction', 'What It Contains'],
    [
        ['Production MongoDB', 'MONGO_URI_PROD',    'Read only',  'Chatbot conversations, Anandita call transcripts, analytics scores, event tracking data'],
        ['Results MongoDB',    'MONGO_URI_RESULTS', 'Write only', 'QA evaluation results, health scores, recommendations, human feedback submissions'],
    ],
    col_widths=[1.6, 1.8, 1.1, 2.1],
)

add_heading(doc, 'The 4-Hour Cycle', 2)
add_body(doc,
    'The runner fires every 4 hours. It processes only conversations and calls that arrived '
    'since the last run — anything already evaluated is skipped via deduplication. '
    'This means the system is always current and never wastes compute re-processing old data. '
    'A single run on a typical day takes 5-15 minutes depending on volume.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 3 — TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════
add_heading(doc, '3.  The Technology Stack', 1)

add_table(doc,
    ['Component', 'Technology', 'Why This Choice'],
    [
        ['Language',         'Python 3.11+',         'Standard, fast to write, excellent MongoDB driver'],
        ['LLM',              'llama3.1:8b via Ollama','Free, offline, no API cost, runs on local Mac. Temperature 0.1 for near-deterministic output.'],
        ['Production DB',    'MongoDB (MONGO_URI_PROD)',   'Read-only connection. Chatbot and voice data already lives here.'],
        ['Results DB',       'MongoDB (MONGO_URI_RESULTS)','Separate cluster. All QA output goes here. Never touches production.'],
        ['Notifications',    'Microsoft Teams via Power Automate webhook', 'ASBL team already uses Teams. Webhook receives a JSON POST and posts to a channel.'],
        ['Scheduling',       'Python time.sleep() loop in runner.py', 'Simple, no external dependency. Runs as a persistent Mac process.'],
        ['Document generation','python-docx',         'Generates this Word document programmatically.'],
    ],
    col_widths=[1.5, 2.0, 3.1],
)

add_heading(doc, 'The LLM — Why Local', 2)
add_body(doc,
    'All AI reasoning in this system uses a local LLM running on the Mac via Ollama. '
    'This was a deliberate choice: the system evaluates real customer conversations which '
    'contain personal details, project enquiries, and pricing discussions. Sending that data '
    'to a cloud API creates a privacy exposure. Running it locally means the data never leaves '
    'the machine.')

add_body(doc,
    'llama3.1:8b is the model. It is capable enough for structured evaluation tasks — '
    'checking facts against a knowledge base, identifying issue types, scoring conversations, '
    'and generating fix descriptions. Temperature is set to 0.1 (near-deterministic) so the '
    'same conversation evaluated twice produces the same result. Max tokens is 1,500 per call.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 4 — CHATBOT QA AGENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '4.  Agent 1 — Chatbot QA', 1)
add_body(doc, 'Files: agents/chatbot_qa.py   |   agents/prompts/chatbot_qa.md', bold=True)

add_heading(doc, 'What It Does', 2)
add_body(doc,
    'The Chatbot QA Agent reads every chatbot conversation from the last 4 hours, '
    'evaluates each one against the knowledge base using the local LLM, and saves a '
    'structured result. If a conversation fails, a Teams notification is fired.')

add_heading(doc, 'Complete Step-by-Step Flow', 2)
add_code(doc,
'1.  Compute window: datetime.now(UTC) - 4 hours\n'
'2.  Query asbl_loft.conversations:\n'
'      createdAt >= window  AND  turnCount >= 2\n'
'      (1-turn conversations have too little content to evaluate)\n'
'3.  Create one MongoClient for the entire batch (shared for efficiency)\n'
'4.  For each conversation:\n'
'      a. already_evaluated("chatbot_qa", "conversation_id", cid)?\n'
'            YES → skip, increment skipped counter\n'
'            NO  → continue\n'
'      b. load_agent_prompt()\n'
'            Opens agents/prompts/chatbot_qa.md fresh from disk\n'
'            (fresh = any edit to the file takes effect on the next run)\n'
'      c. load_knowledge_base()\n'
'            Reads all .md files in knowledge_base/webbot/ sorted alphabetically\n'
'            Joins them with --- separators into one string\n'
'            (8 files: system prompt, project KB, market intelligence, etc.)\n'
'      d. format_conversation(conversationDepth)\n'
'            For each turn:\n'
'              Strip HTML tags from botText\n'
'              If artifactLabel exists: "Bot [showed: unit_plans]: <text>"\n'
'              Otherwise:              "Bot: <text>"\n'
'              Always:                 "User: <text>"\n'
'      e. get_session_signals(client, sessionId)\n'
'            Queries asbl_loft.events where sessionId matches\n'
'            Reads uniqueEventNames array\n'
'            Returns: lead_captured, lead_submitted, visit_booked,\n'
'                     visit_otp_verified, visit_form_started\n'
'      f. Assemble full prompt:\n'
'            [chatbot_qa.md evaluation rules]\n'
'            [8 KB files merged]\n'
'            [SESSION SIGNALS block]\n'
'            [formatted conversation]\n'
'      g. ask_json(prompt) → Ollama returns JSON\n'
'      h. Attach metadata: conversation_id, turn_count, evaluated_at, session_signals\n'
'      i. save_chatbot_result(result) → write to qa_results.chatbot_qa\n'
'      j. If status == FAIL → notify() → Teams webhook\n'
'5.  Print summary: evaluated X, skipped Y'
)

add_heading(doc, 'Why artifactLabel Matters', 2)
add_body(doc,
    'When a user clicks a CTA button on the ASBL website (e.g. "Show me unit plans", '
    '"View pricing", "Book a visit"), that button does two things: it pre-fills the '
    'chat text box with a predefined message, and it tags the resulting turn with an '
    'artifactLabel. This label tells us what the bot actually showed — not just what '
    'it said in text.')

add_table(doc,
    ['artifactLabel', 'What Was Shown to the User'],
    [
        ['unit_plans',   'Floor plan images and unit size breakdown for a project'],
        ['price',        'Pricing card with per-sqft or total price and GST breakdown'],
        ['visit',        'Site visit booking form or call-to-action'],
        ['rental_offer', 'Loft rental income offer card (₹50/sqft/month until Dec 2026)'],
        ['commute',      'Commute time and location map for the project'],
        ['amenity',      'Amenities list for a specific project'],
        ['(empty)',      'Regular typed message — no CTA button was used'],
    ],
    col_widths=[1.6, 5.0],
)

add_body(doc,
    'Exposing this to the LLM evaluator means it can see "Bot [showed: unit_plans]: ..." '
    'and understand that the bot displayed visual floor plan content alongside its text '
    'response. This gives context for evaluating whether the response was complete and appropriate.',
    italic=True)

add_heading(doc, 'Why Session Signals Matter', 2)
add_body(doc,
    'Lead capture on the ASBL website happens via a form overlay — a browser popup that '
    'appears separately from the chat. The form is not part of the conversation transcript. '
    'This means reading the chat text alone cannot tell you whether a lead was captured. '
    'The LLM would have to guess, and it would guess wrong frequently.')

add_body(doc,
    'Instead, the system looks up the real event data from asbl_loft.events using the '
    'sessionId. This collection has 190,306 session documents. Each one contains a '
    'uniqueEventNames array with every distinct event that fired in that session. '
    'The key events are:')

add_table(doc,
    ['Event Name', 'Meaning'],
    [
        ['lead_success',           'Lead form was successfully submitted — lead captured'],
        ['lead_submit',            'Lead form submit button clicked (may or may not have succeeded)'],
        ['visit_booking',          'Site visit booking completed'],
        ['visit_otp_verify_click', 'User clicked verify on the visit OTP screen'],
        ['visit_phone_focus',      'User tapped the phone field on the visit form'],
        ['visit_name_focus',       'User tapped the name field on the visit form'],
    ],
    col_widths=[2.2, 4.4],
)

add_body(doc,
    'These signals are injected into the prompt as a SESSION SIGNALS block above the '
    'conversation. The evaluator is instructed to treat them as authoritative ground truth — '
    'if lead_captured = True, never flag a lead capture failure regardless of what the '
    'conversation text shows.')

add_heading(doc, 'Issue Types and What They Catch', 2)
add_table(doc,
    ['Issue Type', 'Severity', 'What It Catches'],
    [
        ['KB_MISMATCH',
         'HIGH',
         'Any factual claim the bot made that contradicts what the KB says. Covers: price, per-sqft rate, GST rate, GST disclosure, RERA number, handover date, unit size, configuration availability, payment terms, booking amount, offer details, project location. Any numeric discrepancy counts — no benefit of the doubt on numbers.'],
        ['INVENTED_FACT',
         'HIGH',
         'Bot stated something company-specific that has no basis in the KB and could not have come from an external source. Distinct from KB_MISMATCH: this is a fabricated claim, not a wrong version of a real one. Test: is this specific to the company or product? If yes and absent from KB → flag.'],
        ['GUARDRAIL_VIOLATION',
         'HIGH',
         'Bot said something explicitly forbidden: legal possession guarantee, investment return promise, false urgency/scarcity not in KB, negative competitor comparison. Checked from the buyer\'s perspective — soft language like "you have nothing to worry about" or "ASBL always delivers on time" is flagged the same as explicit guarantees.'],
        ['INCOMPLETE_RESPONSE',
         'MEDIUM',
         'User asked something the KB covers and the bot did not fully answer it. Scope is the entire KB — not just prices or RERA. If user asked multiple questions, all must be answered. Only count specific answerable questions, not vague filler words.'],
        ['WRONG_PROJECT_RECOMMENDATION',
         'MEDIUM',
         'Clearly inappropriate recommendation: 30%+ budget mismatch with no acknowledgment, sold-out configuration offered, wrong location when user specified one, or ignored a clearly fitting option. Showing a project modestly above budget is correct sales behaviour — not flagged.'],
        ['TONE_ISSUE',
         'MEDIUM',
         'Unprofessional language for a premium brand: overly casual, dismissive, pushy, high-pressure. Also flagged: deflecting a genuine user concern without addressing it.'],
        ['LANGUAGE_MISMATCH',
         'LOW',
         'User wrote sustained messages in Hindi or Telugu and bot replied only in English without acknowledging.'],
        ['MINOR_PHRASING',
         'LOW',
         'Slightly awkward wording — information is factually correct. Captures quality observations without inflating severity.'],
    ],
    col_widths=[1.9, 0.8, 4.0],
)

add_heading(doc, 'Turn-by-Turn Evaluation Methodology', 2)
add_body(doc,
    'The evaluator does not read the conversation as a single block. It uses a two-pass method:')
add_bullet(doc, 'Pass 1: Read the entire conversation to understand context — what project, budget, and requirements emerged. What language the user used.')
add_bullet(doc, 'Pass 2: Go turn by turn. For each bot response: what did the user ask? Did the bot answer all of it? Is everything said accurate per KB? Was extra information accurate?')
add_body(doc,
    'This prevents the LLM from skipping a mid-conversation turn where a price mismatch '
    'occurs while the overall conversation looks fine. Issues can appear in one turn while '
    'every other turn is clean.',
    italic=True)

add_heading(doc, 'FAIL Condition and Scoring', 2)
add_code(doc,
'FAIL if:  1 or more HIGH severity issues\n'
'      OR  1 or more MEDIUM severity issues\n\n'
'PASS if:  no HIGH issues AND no MEDIUM issues\n\n'
'Score 10: perfect  |  9: one LOW  |  8: LOW issues only\n'
'Score 7: one MEDIUM  |  6: two MEDIUM  |  5: two+ MEDIUM or one HIGH partial recovery\n'
'Score 4: one HIGH  |  3: multiple HIGH  |  2: systematic wrong info  |  1: complete failure'
)

add_heading(doc, 'Confidence Scoring', 2)
add_body(doc,
    'Every evaluation includes a confidence score (0.0–1.0) representing how certain the '
    'LLM is about its own assessment. This is separate from the quality score. It feeds '
    'directly into the health score formula — a speculative flag with confidence 0.4 hurts '
    'the health score less than a confirmed flag with confidence 0.95.')

add_table(doc,
    ['Confidence', 'Meaning', 'Example'],
    [
        ['1.0', 'Exact mismatch verified against KB. No ambiguity.',          'KB says 1.94 cr, bot said 1.90 cr. Certain.'],
        ['0.8', 'Strong evidence, minor conversational ambiguity.',             'Price likely wrong but unclear which unit size was meant.'],
        ['0.6', 'Probable issue, short conversation or ambiguous question.',   'Incomplete answer but only 2 turns total.'],
        ['0.4', 'Possible issue. Genuinely uncertain.',                        'Tone call where reasonable people might disagree.'],
        ['0.2', 'Very uncertain. Flag is speculative.',                        'Something seems off but no clear evidence.'],
    ],
    col_widths=[1.0, 2.2, 3.4],
)
add_body(doc,
    'Calibration rule: MEDIUM severity issues should sit at 0.6–0.8, not cluster at 0.9+. '
    'Reserve 0.9+ for confirmed factual mismatches only. 1–2 turn conversations cap at 0.5.',
    italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 5 — VOICE QA AGENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '5.  Agent 2 — Voice QA (Anandita)', 1)
add_body(doc, 'Files: agents/voice_qa.py   |   agents/prompts/voice_qa.md', bold=True)

add_heading(doc, 'What It Does', 2)
add_body(doc,
    'The Voice QA Agent reads every Anandita phone call transcript from the last 4 hours '
    'and evaluates it for factual accuracy, decimal price format, language handling, and tone. '
    'Unlike the chatbot, voice calls have unique failure modes — a price spoken incorrectly '
    'sounds like a completely different number to the caller, and a language mismatch on a '
    'call causes the caller to disengage immediately.')

add_body(doc,
    'The agent produces two independent outputs for each call: a quality score (same 1-10 '
    'rubric as chatbot) and a separate language_compliance field (PASS/FAIL). A call can '
    'score 8/10 and still have language_compliance = FAIL if Anandita missed a language switch.')

add_heading(doc, 'Complete Step-by-Step Flow', 2)
add_code(doc,
'1.  Compute window: datetime.now(UTC) - 4 hours\n'
'2.  Query ASBLVoiceBot.call_transcripts:\n'
'      started_at >= window  AND  transcript is not empty\n'
'3.  For each call:\n'
'      a. already_evaluated("voice_qa", "call_sid", sid)?\n'
'            YES → skip\n'
'            NO  → continue\n'
'      b. load_agent_prompt() → voice_qa.md fresh from disk\n'
'      c. Load 3 KB files from knowledge_base/anandita/:\n'
'            project_facts.md, system_prompt.md, qa_checklist.md\n'
'      d. format_transcript(transcript)\n'
'            speaker == "Anandita"    → "Anandita: <text>"\n'
'            speaker == phone number  → "Caller: <text>"\n'
'      e. Attach call metadata to prompt:\n'
'            call_sid, phone_number, call_direction, language_used, project\n'
'      f. ask_json(prompt) → Ollama returns JSON\n'
'      g. Save result to qa_results.voice_qa\n'
'      h. If status == FAIL → notify() → Teams webhook'
)

add_heading(doc, 'Issue Types', 2)
add_body(doc,
    'Voice uses all the same issue types as chatbot plus two voice-specific ones (DECIMAL_NUMBER and LANGUAGE_HANDLING). '
    'All nine are listed below:')

add_table(doc,
    ['Issue Type', 'Severity', 'Detail'],
    [
        ['KB_MISMATCH',
         'HIGH',
         'Anandita stated a fact that directly contradicts the KB. Covers price, per-sqft rate, GST disclosure, RERA number, handover date, unit size, configuration availability, payment terms, booking amount, offer details, project location. Any numeric discrepancy counts — no benefit of the doubt on numbers.'],
        ['INVENTED_FACT',
         'HIGH',
         'Anandita stated something company-specific that has no basis in the KB and could not have come from an external source. KB_MISMATCH = KB has the fact but she got it wrong. INVENTED_FACT = the KB does not have this at all and it is specific enough to the company that it must be grounded in the KB.'],
        ['GUARDRAIL_VIOLATION',
         'HIGH',
         'Anandita said something explicitly forbidden: legal possession guarantee, investment return promise, false urgency/scarcity not in KB, negative competitor comparison, AI self-identification ("As an AI language model..."), or "I\'ll call you back" (she is the call). Checked from the buyer\'s perspective — soft language like "ASBL always delivers on time" is flagged the same as explicit guarantees.'],
        ['DECIMAL_NUMBER',
         'HIGH',
         'Anandita spoke a decimal price in the wrong verbal format. "One ninety-four crore" sounds like 194 crore to the caller — a completely different number. Correct: "one point nine four crore". This is the most common and most damaging voice-specific error. DECIMAL_NUMBER = right number, wrong format. KB_MISMATCH = wrong number. Both can fire on the same price statement.'],
        ['INCOMPLETE_RESPONSE',
         'MEDIUM',
         'Caller asked something the KB covers and Anandita did not fully answer it. Scope is the entire KB — not just prices or RERA. If the caller asked multiple questions, all must be addressed.'],
        ['WRONG_PROJECT_RECOMMENDATION',
         'MEDIUM',
         'Clearly inappropriate recommendation: 30%+ budget mismatch with no acknowledgment, sold-out configuration offered, wrong location when caller specified one, or ignored a clearly fitting option. Showing a project modestly above budget is correct sales behaviour — not flagged.'],
        ['TONE_ISSUE',
         'MEDIUM',
         'Unprofessional language for a premium brand: overly casual, dismissive, pushy, high-pressure. Also flagged: deflecting a genuine caller concern without addressing it.'],
        ['LANGUAGE_HANDLING',
         'MEDIUM',
         'Caller spoke sustained Hindi or Telugu and Anandita stayed in English. On a voice call this is more serious than in text — the caller feels unheard and is more likely to disengage. Sustained sentences = MEDIUM. Occasional mixed words = LOW, note only.'],
        ['MINOR_PHRASING',
         'LOW',
         'Slightly awkward wording or transcription-adjacent phrasing. Information is factually correct. Captures quality observations without inflating severity.'],
    ],
    col_widths=[1.9, 0.8, 3.9],
)

add_heading(doc, 'Language Compliance — Separate Field', 2)
add_body(doc,
    'Language compliance is evaluated independently from the quality score and stored '
    'as its own field in the result document.')
add_code(doc,
'language_compliance = PASS  if:\n'
'  Entire call was in English and caller spoke English\n'
'  OR caller used Hindi/Telugu and Anandita switched to match and maintained it\n\n'
'language_compliance = FAIL  if:\n'
'  Caller clearly spoke Hindi or Telugu (multiple sentences)\n'
'  AND Anandita stayed in English throughout\n'
'  OR Anandita switched appropriately but then reverted without caller switching back'
)

add_heading(doc, 'GUARDRAIL — Voice-Specific Additions', 2)
add_body(doc,
    'In addition to the shared guardrail rules (guarantees, return promises, false urgency), '
    'voice has two specific forbidden behaviours:')
add_bullet(doc, 'Exposing AI identity — "As an AI language model...", "I am an AI assistant", "I don\'t have real-time data". Anandita has a persona. Breaking it destroys trust.')
add_bullet(doc, 'Saying "I\'ll call you back" — Anandita IS the call. She cannot schedule herself to call anyone. This is factually wrong and must be flagged.')

add_heading(doc, 'Transcription Artifact Rule', 2)
add_body(doc,
    'Voice evaluations are based on speech-to-text transcripts, which are imperfect. '
    'A single garbled word or missing turn due to poor audio is a transcription artifact, '
    'not a quality failure. The evaluator is instructed: flag patterns that repeat across '
    'multiple lines, not isolated noise. If the call has obvious gaps or repeated [inaudible] '
    'markers, confidence is reduced accordingly.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 6 — ANALYTICS AGENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '6.  Agent 3 — Analytics Validator', 1)
add_body(doc, 'File: agents/analytics.py   (no LLM — pure data checks)', bold=True)

add_heading(doc, 'What It Does and Why It Is Different', 2)
add_body(doc,
    'The Analytics Agent is fundamentally different from the QA agents. It does not evaluate '
    'natural language conversations. It runs four rule-based data checks directly against '
    'the production analytics database and produces structured flags. No LLM is involved. '
    'The flags go into the results database and are picked up by the Feedback Agent.')

add_body(doc,
    'The purpose is to validate the analytics scoring system itself — the external system '
    'that assigns bands (Band1_Spark through Band5_Hot) to leads and fires Meta/Google '
    'conversion events when a lead crosses a threshold. If the scoring system is '
    'miscalibrated, it will waste Meta ad budget by firing expensive events for leads '
    'who never actually convert.')

add_heading(doc, 'Check 1 — Band Conversion Rates', 2)
add_body(doc,
    'For every active outcome defined in config/outcomes.json, the agent computes the '
    'actual conversion rate per scoring band and compares it to the target you set.')
add_code(doc,
'For each active outcome (e.g. site_visit_booked):\n'
'  For each lead in analytics_db.scores_overall:\n'
'    band      = lead.lifetime.highest_band_ever\n'
'    converted = check outcome field on this lead\n'
'\n'
'  For each band (Band1_Spark through Band5_Hot):\n'
'    rate        = converted_leads / total_leads_in_band\n'
'    target      = target % from outcomes.json for this band\n'
'    meta_events = count of meta_conversion_events fired for this band\n'
'\n'
'    FLAG (HIGH — BAND_OUTCOME_MISMATCH) if:\n'
'      meta_events > 0     (money is being spent on this band)\n'
'      AND rate < target   (conversions are below your target)\n'
'      AND total >= 3      (enough data to judge)'
)

add_heading(doc, 'Check 2 — Multiplier Effectiveness', 2)
add_body(doc,
    'Multipliers are score boosts applied when a lead takes a high-intent action. '
    'This check verifies that each multiplier is actually correlated with better outcomes. '
    'If completing multiplier M1 does not improve visit booking rates, the multiplier '
    'is boosting scores without improving lead quality.')
add_table(doc,
    ['Multiplier', 'What Triggers It', 'Effect'],
    [
        ['M1', 'Affordability calculator (EMI-based) determines the lead CAN afford the property', 'Score boost'],
        ['M2', 'Lead stayed on site 30+ seconds after seeing they CANNOT afford it — shows commitment', 'Score boost'],
        ['M3', 'Lead typed their actual home address — high intent signal', 'Score boost'],
        ['M4', 'Lead completed OTP phone verification', '1.3× score multiplier'],
    ],
    col_widths=[0.6, 4.2, 1.8],
)
add_code(doc,
'For each multiplier:\n'
'  Group A = leads who completed this multiplier\n'
'  Group B = leads who did NOT complete this multiplier\n'
'  rate_A  = visit_booked count in A / total in A\n'
'  rate_B  = visit_booked count in B / total in B\n'
'\n'
'  FLAG (MEDIUM — MULTIPLIER_INEFFECTIVE) if:\n'
'    rate_A <= rate_B   (multiplier not improving outcomes)\n'
'    AND sample_size >= 5'
)

add_heading(doc, 'Check 3 — Meta Signal Accuracy', 2)
add_body(doc,
    'Every time a lead crosses a band threshold, a Meta conversion event fires. '
    'Each event has a monetary value assigned to it — Band5_Hot fires an event worth '
    '₹7,85,500. This check measures what fraction of those fired events actually result '
    'in a site visit booking.')
add_code(doc,
'For each meta_conversion_event fired:\n'
'  Find the lead (via lead_id, or session_id → scores_session_wise → lead_id)\n'
'  Check: has this lead booked a site visit?\n'
'    YES → led_to_visit += 1\n'
'    NO  → wasted += 1\n'
'\n'
'accuracy = led_to_visit / total_events\n'
'\n'
'FLAG (HIGH — META_SIGNAL_LOW_ACCURACY) if:\n'
'  accuracy < 20%  AND  total_events >= 10'
)
add_notice(doc,
    'Current finding from real data: 0.38% accuracy. Only 52 out of 13,653 Meta conversion '
    'events led to a confirmed site visit. Band5_Hot fires ₹7,85,500 per lead. '
    'Approximately 1 in 263 Band5_Hot leads actually converts.',
    'E8691E')

add_heading(doc, 'Check 4 — Funnel Gaps', 2)
add_table(doc,
    ['Gap', 'Query', 'Severity', 'What It Means'],
    [
        ['leads_with_no_session',       'scores_overall where lifetime.total_sessions == 0',                           'LOW',    'Leads scored but never had a website session recorded — possible data integrity issue'],
        ['affordability_yes_no_visit',  'scores_overall where affordability_outcome = "YES" AND has_visit_booked = False', 'MEDIUM', 'Leads confirmed they can afford but never booked — ops follow-up failure, not AI failure'],
    ],
    col_widths=[1.8, 2.4, 0.8, 1.6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 7 — FEEDBACK AGENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '7.  Agent 4 — Feedback Agent', 1)
add_body(doc, 'Files: agents/feedback.py   |   agents/prompts/feedback_reasoning.md', bold=True)

add_heading(doc, 'What It Does', 2)
add_body(doc,
    'The Feedback Agent is the manager of the system. It reads all outputs from all three '
    'QA agents, computes health scores using a weighted penalty formula, and then sends '
    'the entire picture to the LLM to identify every genuine problem — including patterns '
    'that span multiple agents and that no single-agent check can find.')

add_notice(doc,
    'Key design principle: the LLM sees everything at once — all chatbot results, all voice '
    'results, all analytics flags, and the full health score history — and reasons about the '
    'complete picture. There are no hardcoded thresholds for problem detection. '
    'The LLM decides what is a problem based on evidence.',
    '1A376C')

add_heading(doc, 'Phase 1 — Health Score Computation (No LLM)', 2)
add_body(doc, 'Health scores are computed separately for chatbot, voice, and analytics.')

add_heading(doc, 'Chatbot and Voice Health Score Formula', 3)
add_code(doc,
'For each evaluation result this cycle:\n'
'\n'
'  base_score  = Ollama raw quality score (1–10)\n'
'  confidence  = Ollama confidence (0.0–1.0, default 0.5 if missing)\n'
'\n'
'  penalty = 0\n'
'  severity_weight = { HIGH: 2.0, MEDIUM: 0.8, LOW: 0.2 }\n'
'  for each issue in result.issues:\n'
'    penalty += severity_weight[severity] × confidence\n'
'\n'
'  adjusted_score = clamp(base_score - penalty, 0, 10)\n'
'\n'
'health_score = average(all adjusted_scores this cycle)\n'
'fail_rate    = count(FAIL results) / total results this cycle'
)

add_body(doc,
    'Confidence is included in the penalty formula deliberately. A HIGH issue that the LLM '
    'is 95% confident about should penalise the health score more than a HIGH issue the LLM '
    'is 40% confident about. This prevents speculative flags from unfairly degrading the '
    'health score as much as confirmed findings.')

add_heading(doc, 'Analytics Health Score Formula', 3)
add_code(doc,
'analytics_score = max(0, 10 - (high_flags × 2.5) - (medium_flags × 1.0))'
)

add_heading(doc, 'Trend Detection', 3)
add_code(doc,
'Fetch last 3 health score records for this component:\n'
'  All three decreasing (each < previous) → trend = DECLINING\n'
'  All three increasing (each > previous) → trend = IMPROVING\n'
'  Otherwise                              → trend = STABLE'
)

add_heading(doc, 'Health Score Scale', 3)
add_table(doc,
    ['Score Range', 'Status', 'Meaning'],
    [
        ['8.0 – 10.0', 'Healthy',    'System performing well this cycle'],
        ['6.5 – 7.9',  'Acceptable', 'Minor issues, nothing alarming'],
        ['5.0 – 6.4',  'Concerning', 'Noticeable quality problems — review this cycle\'s failures'],
        ['3.0 – 4.9',  'Bad',        'Significant failures — action required soon'],
        ['0.0 – 2.9',  'Critical',   'Major breakdown — immediate action required'],
    ],
    col_widths=[1.4, 1.2, 4.0],
)

add_heading(doc, 'Phase 2 — LLM Problem Detection', 2)
add_body(doc,
    'Once health scores are saved, the full picture is assembled into one prompt and sent '
    'to Ollama. The LLM receives the feedback_reasoning.md instructions plus all QA results, '
    'all analytics flags, and the full health score history for this component.')

add_body(doc, 'Cross-agent patterns the LLM is taught to identify:', bold=True)
add_table(doc,
    ['Pattern Name', 'Condition', 'Root Cause Diagnosis'],
    [
        ['Analytics miscalibrated',  'Chatbot ≥6.5 AND voice ≥6.5, but outcomes below target',   'Both bots are performing well. The scoring system is promoting wrong leads.'],
        ['Shared KB wrong',          'Same HIGH issue type in both chatbot AND voice results',     'The knowledge base has an error — it is not a bot-specific bug.'],
        ['Band thresholds too low',  'Meta accuracy <5% AND QA scores are healthy',               'Leads reach Band5_Hot without real purchase intent. Thresholds need raising.'],
        ['Doubly broken',            'Meta accuracy <5% AND QA scores also failing',              'Two separate root causes at once — fix both independently.'],
        ['Voice-specific bug',       'Chatbot score > voice score by 3+ points',                  'Anandita-specific issue, not shared KB.'],
        ['Chatbot-specific bug',     'Voice score > chatbot score by 3+ points',                  'Website chatbot-specific issue, not shared KB.'],
        ['Affordability ops gap',    'affordability_yes_no_visit count > 10',                     'Sales ops failing to follow up on confirmed-affordable leads — not an AI failure.'],
        ['Multiplier miscalibrated', 'Multiplier ineffective AND QA scores healthy',              'Scoring weight too high for this signal. Miscalibration, not a bot bug.'],
        ['DECLINING trend',          'Health score falling for 3+ consecutive cycles',            'Something changed recently — check git commits and KB updates around that date.'],
    ],
    col_widths=[1.8, 2.1, 2.7],
)

add_heading(doc, 'Problem Object Structure', 2)
add_code(doc,
'{\n'
'  "id":               "P001",\n'
'  "title":            "Short description of the problem",\n'
'  "urgency":          "HIGH",         // HIGH / MEDIUM / LOW\n'
'  "type":             "SYSTEMATIC_BUG",\n'
'  "components":       ["voice"],\n'
'  "description":      "Plain English explanation with numbers",\n'
'  "evidence":         ["specific data point 1", "data point 2"],\n'
'  "what_is_wrong":    "The specific broken thing — one line",\n'
'  "what_is_not_wrong":"What should NOT be blamed — one line"\n'
'}'
)
add_table(doc,
    ['Problem Type', 'When It Applies'],
    [
        ['SYSTEMATIC_BUG', 'Same issue appearing in >50% of evaluations this cycle'],
        ['CALIBRATION',    'Analytics scoring model is miscalibrated'],
        ['KB_OUTDATED',    'Knowledge base data is wrong or stale'],
        ['PROCESS_GAP',    'Ops or human process is broken — not an AI issue'],
        ['TREND',          'Health score declining consistently across multiple cycles'],
    ],
    col_widths=[1.8, 4.8],
)

add_heading(doc, 'Human Feedback Integration', 2)
add_body(doc,
    'The Feedback Agent also accepts manually submitted feedback from the team. '
    'Submissions are stored in qa_results.feedback and included in the LLM\'s context '
    'when looking for problems. This allows human observations to be factored into '
    'the automated analysis.')
add_code(doc,
'# Submit feedback\n'
'python3 agents/feedback.py submit "Bot gave wrong Spectra price" "Anirudh"\n'
'\n'
'# View recent submissions\n'
'python3 agents/feedback.py list 20'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 8 — RECOMMENDATION AGENT
# ════════════════════════════════════════════════════════════════
add_heading(doc, '8.  Agent 5 — Recommendation Agent', 1)
add_body(doc, 'Files: agents/recommendation.py   |   agents/prompts/recommendation.md', bold=True)

add_heading(doc, 'What It Does', 2)
add_body(doc,
    'The Recommendation Agent receives the problem list directly from the Feedback Agent '
    'in memory — no database read between agents. For each problem it produces one specific, '
    'actionable fix: what to change, exactly where to change it, and what will improve once '
    'the fix is made. It is the engineer in the pipeline.')

add_heading(doc, 'Fix Generation Logic', 2)
add_code(doc,
'Receive problems list from Feedback Agent (in memory, not re-fetched)\n'
'Sort: HIGH urgency first, then MEDIUM, then LOW\n'
'\n'
'For each problem:\n'
'  Step 1: Check KNOWN_FIXES dictionary\n'
'    DECIMAL_NUMBER    → template fix applied directly (no LLM needed)\n'
'    PRICE_ACCURACY    → template fix\n'
'    RERA_NUMBER       → template fix\n'
'    LANGUAGE_HANDLING → template fix\n'
'\n'
'  Step 2: No template match → call Ollama with:\n'
'    [recommendation.md — ASBL file locations, project facts, band/multiplier context]\n'
'    [the problem: id, title, components, evidence, what_is_wrong]\n'
'    [historical baseline: first score + date, best score + date, current score + date]\n'
'\n'
'  Step 3: Save fix object to results database'
)

add_heading(doc, 'Historical Baseline in Fixes', 2)
add_body(doc,
    'For each component (chatbot, voice, analytics), the agent fetches the last 30 days '
    'of health score history from the results database. This is given to the LLM so it can '
    'identify when a problem started — making the fix more precise.')
add_code(doc,
'Baseline example given to LLM:\n'
'  chatbot: first score 7.8 on 2026-05-01, best 8.2 on 2026-05-05, current 4.1 on 2026-05-18\n'
'\n'
'LLM output:\n'
'  "Voice score was 7.5 on 2026-05-09 and dropped to 2.5 by 2026-05-12.\n'
'   Something likely changed around 2026-05-10 — check KB edits and deployments."'
)

add_heading(doc, 'Fix Object Structure', 2)
add_code(doc,
'{\n'
'  "problem_id":       "P001",\n'
'  "problem_title":    "Anandita decimal format broken",\n'
'  "urgency":          "HIGH",\n'
'  "root_cause":       "One sentence — includes when it started from historical baseline",\n'
'  "fix":              "Exactly what to change — specific enough to act on today",\n'
'  "where":            "Exact file path, or scoring model, or CRM/ops workflow",\n'
'  "change_type":      "kb_file_edit",\n'
'  "expected_outcome": "What improves, by how much, how quickly"\n'
'}'
)
add_table(doc,
    ['change_type', 'Meaning', 'Takes Effect'],
    [
        ['kb_file_edit',  'Edit a .md file in knowledge_base/',      'Next 4-hour run — automatic'],
        ['config_edit',   'Edit a .json file in config/',             'Next 4-hour run — automatic'],
        ['scoring_model', 'Change in the analytics scoring system',   'Requires deployment — 24–48 hours'],
        ['ops_process',   'Human or CRM action needed — no code',    'When a human acts on it'],
    ],
    col_widths=[1.5, 2.8, 2.3],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 9 — THE RUNNER
# ════════════════════════════════════════════════════════════════
add_heading(doc, '9.  The Runner — How Everything Connects', 1)
add_body(doc, 'File: runner.py', bold=True)

add_body(doc,
    'The runner is the orchestrator. It runs on an infinite loop with a 4-hour sleep '
    'between cycles. Each cycle calls all 5 agents in the correct sequence. '
    'The problems list from the Feedback Agent is passed directly to the '
    'Recommendation Agent in memory — no database round-trip between those two steps.')

add_code(doc,
'def run_once():\n'
'    chatbot_qa.batch(hours=4)        # Step 1\n'
'    voice_qa.batch(hours=4)          # Step 2\n'
'    analytics.run()                  # Step 3\n'
'    problems = feedback.aggregate()  # Step 4 — returns list of Problem objects\n'
'    recommendation.run(problems)     # Step 5 — receives list, no re-fetch\n'
'\n'
'while True:\n'
'    run_once()\n'
'    time.sleep(4 * 3600)             # 4 hours\n'
'\n'
'# python3 runner.py once → run one cycle now (for testing)\n'
'# python3 runner.py      → run indefinitely every 4 hours (production)'
)

add_heading(doc, 'Deduplication — How It Works', 2)
add_body(doc,
    'Each agent checks whether a given item has already been evaluated before processing it. '
    'This is handled by the already_evaluated() function in db.py which queries the results '
    'database for a document matching the unique identifier.')
add_code(doc,
'already_evaluated("chatbot_qa", "conversation_id", cid)\n'
'  → queries qa_results.chatbot_qa for {conversation_id: cid}\n'
'  → returns True if found (skip), False if not found (process)\n'
'\n'
'Same pattern for voice_qa using call_sid as the unique identifier.'
)
add_body(doc,
    'This means if the system crashes mid-cycle, conversations already evaluated in that '
    'cycle are not re-evaluated on the next run. The window query (last 4 hours) ensures '
    'that genuinely new conversations are always picked up.',
    italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 10 — PROMPT DESIGN PHILOSOPHY
# ════════════════════════════════════════════════════════════════
add_heading(doc, '10.  The Prompt Design Philosophy', 1)

add_body(doc,
    'The prompt files define how to evaluate. The knowledge base defines what to evaluate '
    'against. These are two separate concerns and they are kept separate deliberately.')

add_heading(doc, 'Why No Hardcoded Facts in Prompts', 2)
add_body(doc,
    'Earlier versions of the chatbot QA prompt contained hardcoded prices, RERA numbers, '
    'project names, offer details, and specific examples using real figures. '
    'This created a maintenance problem: whenever a price changed or a new project launched, '
    'both the knowledge base AND the prompt had to be updated. If only one was updated, '
    'the system would evaluate against stale data and produce wrong flags.')

add_body(doc,
    'The current design removes all facts from the prompts. The prompts contain only '
    'evaluation logic — rules, reasoning instructions, issue type definitions, and output '
    'format. All facts come from the KB injected at runtime. '
    'Update the KB once, the prompt works correctly automatically.')

add_heading(doc, 'Generic Rules Over Specific Examples', 2)
add_body(doc,
    'Issue type definitions are written as general rules that apply to anything in the KB, '
    'not as specific checks for a list of named fields. For example, INCOMPLETE_RESPONSE '
    'used to list specific fields (price, RERA, handover date, unit size). It now says: '
    '"any topic covered in the KB". This means the same prompt works correctly whether the '
    'user asked about payment plans, amenities, legal registration, or rental offers — '
    'without any prompt change.')

add_heading(doc, 'Prompt Files and When They Take Effect', 2)
add_body(doc,
    'Every agent that uses the LLM loads its prompt file fresh from disk at the start of '
    'each run. There is no caching. Edit the file and the next 4-hour cycle picks up the '
    'change automatically. No code deployment, no restart.')

add_table(doc,
    ['Prompt File', 'Size', 'Controls'],
    [
        ['agents/prompts/chatbot_qa.md',        '~6KB', 'Issue types, severity rules, turn-by-turn methodology, scoring rubric, confidence calibration, edge cases, output format'],
        ['agents/prompts/voice_qa.md',           '~6KB', 'Same structure as chatbot plus DECIMAL_NUMBER, LANGUAGE_HANDLING, language compliance evaluation, transcription artifact handling'],
        ['agents/prompts/feedback_reasoning.md', '~12KB','Health score interpretation, cross-agent reasoning patterns, urgency rules, false-positive avoidance, problem type definitions'],
        ['agents/prompts/recommendation.md',     '~13KB','All ASBL file locations, project facts for fixes, band/multiplier context, known fix templates, change type definitions'],
    ],
    col_widths=[2.6, 0.6, 3.4],
)

add_heading(doc, 'Three Key Improvements Applied to Both QA Prompts', 2)
add_table(doc,
    ['Improvement', 'What It Does', 'Why It Was Needed'],
    [
        ['Context loss protection',
         'If KB content cannot be located to verify a claim, lower confidence to 0.5 and note in summary. Do not guess.',
         'The LLM has a finite context window. With 8 KB files + long conversations, early KB content can fade. A missed issue is better than a wrong flag.'],
        ['GUARDRAIL buyer perspective',
         'Apply from the buyer\'s perspective — if a buyer would read it as a commitment, flag it regardless of exact words used.',
         'The LLM was pattern-matching on words like "guaranteed". Soft language like "you have nothing to worry about" carries the same implied commitment.'],
        ['Confidence calibration anchors',
         'Added a self-check question and explicit rule: MEDIUM issues sit at 0.6-0.8, 0.9+ is reserved for confirmed factual mismatches.',
         'Without guidance the LLM clustered at 0.85+ on everything. Confidence became a flat constant and stopped differentiating certain findings from speculative ones.'],
    ],
    col_widths=[1.8, 2.4, 2.4],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 11 — KNOWLEDGE BASE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '11.  The Knowledge Base', 1)

add_body(doc,
    'The knowledge base is the ground truth for all factual evaluation. Both the chatbot '
    'and voice QA agents load the relevant KB files at the start of each evaluation and '
    'inject them into the prompt. The LLM compares every bot claim against this content.')

add_heading(doc, 'Website Chatbot KB (8 files)', 2)
add_table(doc,
    ['File', 'Contents'],
    [
        ['00_qa_checklist.md',          'Evaluation criteria and quality standards for the chatbot'],
        ['01_system_prompt.md',         'The chatbot\'s actual system prompt — its persona, role, and behavioural rules'],
        ['02_kb_project_loft.md',       'ASBL Loft: prices, RERA, handover, rental offer, payment terms, model flat status'],
        ['03_kb_market_intelligence.md','Hyderabad real estate market context, area comparisons, price trends'],
        ['04_kb_competitive_landscape.md','Competitor landscape — used to evaluate how the bot handles competitor questions'],
        ['05_kb_persona_playbook.md',   'How the chatbot should engage different buyer personas'],
        ['06_kb_objection_library.md',  'Standard objections buyers raise and how to handle them'],
        ['07_kb_resale_framework.md',   'Resale market context and investment framing'],
        ['08_deployment_guide.md',      'Technical deployment and operational notes'],
    ],
    col_widths=[2.4, 4.2],
)

add_heading(doc, 'Anandita Voice KB (3 files)', 2)
add_table(doc,
    ['File', 'Contents'],
    [
        ['project_facts.md', 'Prices, RERA numbers, handover dates, configurations for all 4 projects. Primary factual reference for voice evaluation.'],
        ['system_prompt.md', 'Anandita\'s persona definition, language rules, tone guidelines, what she must and must not say'],
        ['qa_checklist.md',  'Voice-specific evaluation criteria including decimal format rules and language compliance expectations'],
    ],
    col_widths=[1.8, 4.8],
)

add_notice(doc,
    'Important: the KB is the single source of truth. The QA prompts contain no facts. '
    'If a price changes, update the relevant KB file. The change takes effect on the next run. '
    'The prompts do not need to be touched.',
    '1A376C')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 12 — PRODUCTION DATABASE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '12.  Production Database Reference', 1)
add_body(doc, 'Connection: MONGO_URI_PROD   |   Access: Read-only. Never written to.', bold=True)

add_heading(doc, 'asbl_loft.conversations — Chatbot conversations (61,155 total)', 2)
add_table(doc,
    ['Field', 'Type', 'Notes'],
    [
        ['conversationId',    'String',  'Unique ID e.g. "c-1778876809751-13ps22ku". Used for deduplication.'],
        ['sessionId',         'String',  'Links to asbl_loft.events for session signal lookup.'],
        ['createdAt',         'Mixed',   'May be ISO string or datetime object — query handles both formats.'],
        ['turnCount',         'Integer', 'Number of user-bot exchanges. Minimum 2 required for evaluation.'],
        ['conversationDepth', 'Array',   'Each turn: {turnNumber, userText, botText (HTML — stripped by formatter), artifactLabel}'],
        ['leadId',            'String',  'May be null. Anonymous visitors are expected and normal.'],
    ],
    col_widths=[1.7, 0.8, 4.1],
)

add_heading(doc, 'asbl_loft.events — Session event tracking (190,306 sessions)', 2)
add_body(doc,
    'Each document represents one website session. '
    'The uniqueEventNames array contains every distinct event that fired in that session.')
add_table(doc,
    ['Event Name', 'What It Confirms'],
    [
        ['lead_success',           'Lead form fully submitted and accepted — authoritative lead capture signal'],
        ['lead_submit',            'Submit button clicked — may or may not have succeeded'],
        ['visit_booking',          'Site visit booking completed'],
        ['visit_otp_verify_click', 'User clicked OTP verify on visit screen'],
        ['visit_phone_focus',      'User tapped phone field on visit form — started the process'],
        ['visit_name_focus',       'User tapped name field on visit form'],
        ['visit_otp_send_click',   'User sent OTP on visit form'],
    ],
    col_widths=[2.2, 4.4],
)

add_heading(doc, 'ASBLVoiceBot.call_transcripts — Anandita calls', 2)
add_table(doc,
    ['Field', 'Type', 'Notes'],
    [
        ['call_sid',          'String',  'Unique call ID. Used for deduplication.'],
        ['phone_number',      'String',  'Caller\'s phone number'],
        ['call_direction',    'String',  '"inbound" or "outbound"'],
        ['started_at',        'DateTime','Call start timestamp — used for 4-hour window query'],
        ['transcript',        'Array',   'Each turn: {speaker, text, ts}. speaker = "Anandita" or caller\'s number.'],
        ['language_used',     'String',  'Language detected on the call'],
        ['project',           'String',  'Which ASBL project was discussed'],
        ['call_outcome',      'String',  'Call result — e.g. NO_ANSWER, COMPLETED. Not yet used in evaluation (pending improvement).'],
        ['site_visit_agreed', 'Boolean', 'Did the caller agree to a site visit? Not yet used in evaluation.'],
    ],
    col_widths=[1.7, 0.8, 4.1],
)

add_heading(doc, 'analytics_db collections', 2)
add_table(doc,
    ['Collection', 'Used For', 'Key Fields'],
    [
        ['scores_overall',            'Band conversion and multiplier checks', 'lead_id, lifetime.highest_band_ever, milestones.has_visit_booked, milestones.has_otp_verified, milestones.affordability_outcome'],
        ['meta_conversion_events',    'Meta signal accuracy check',            'lead_id, session_id, band.name'],
        ['multiplier_completion_events','Multiplier effectiveness check',       'lead_id, pattern_name (M1/M2/M3/M4)'],
        ['scores_session_wise',       'Meta event lead lookup fallback',       'session_id, lead_id'],
    ],
    col_widths=[2.0, 1.8, 2.8],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 13 — RESULTS DATABASE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '13.  Results Database — Full Schema', 1)
add_body(doc, 'Database: qa_results   |   Connection: MONGO_URI_RESULTS', bold=True)

add_heading(doc, 'Collection: chatbot_qa', 2)
add_body(doc, 'One document per chatbot conversation evaluated.')
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['_id',             'ObjectId', 'Auto-generated'],
        ['conversation_id', 'String',   'Unique conversation ID. Used for deduplication on subsequent runs.'],
        ['status',          'String',   '"PASS", "FAIL", or "SKIPPED"'],
        ['score',           'Integer',  '1–10. Ollama raw quality score.'],
        ['confidence',      'Float',    '0.0–1.0. Ollama certainty. Default 0.5 if LLM omits it.'],
        ['turn_count',      'Integer',  'Number of user-bot exchanges in the conversation.'],
        ['issues',          'Array',    'Issue objects: {type, severity (HIGH/MEDIUM/LOW), detail (quoted wrong line + correct value)}'],
        ['summary',         'String',   'One-sentence evaluation summary from Ollama.'],
        ['evaluated_at',    'String',   'ISO timestamp.'],
        ['session_signals', 'Object',   '{lead_captured, lead_submitted, visit_booked, visit_otp_verified, visit_form_started} — from asbl_loft.events'],
    ],
    col_widths=[1.5, 0.9, 4.2],
)

add_heading(doc, 'Collection: voice_qa', 2)
add_body(doc, 'One document per call evaluated.')
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['_id',                'ObjectId', 'Auto-generated'],
        ['call_sid',           'String',   'Unique call ID. Used for deduplication.'],
        ['phone_number',       'String',   'Caller\'s number'],
        ['call_direction',     'String',   '"inbound" or "outbound"'],
        ['language_used',      'String',   'Language detected'],
        ['project',            'String',   'Project discussed'],
        ['status',             'String',   '"PASS", "FAIL", or "SKIPPED"'],
        ['score',              'Integer',  '1–10'],
        ['confidence',         'Float',    '0.0–1.0'],
        ['language_compliance','String',   '"PASS" or "FAIL" — evaluated independently from score'],
        ['issues',             'Array',    'Issue objects: {type, severity, detail}'],
        ['summary',            'String',   'One-sentence summary'],
        ['evaluated_at',       'String',   'ISO timestamp'],
    ],
    col_widths=[1.7, 0.9, 4.0],
)

add_heading(doc, 'Collection: analytics_runs', 2)
add_body(doc, 'One document per analytics cycle.')
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['run_at',                  'String', 'ISO timestamp'],
        ['outcomes_checked',        'Array',  'Names of active outcomes e.g. ["site_visit_booked", "otp_verified"]'],
        ['band_conversion_rates',   'Object', 'Nested: outcome → band → {total_leads, converted, rate, target_pct, meta_events_fired, flag}'],
        ['multiplier_effectiveness','Object', 'Per multiplier M1–M4: {with_rate, without_rate, effective, sample_size}'],
        ['meta_signal_accuracy',    'Object', '{total_events_fired, events_led_to_visit, wasted_signal_count, accuracy_rate}'],
        ['funnel_gaps',             'Array',  'Gap objects: {gap, count, severity, detail}'],
        ['flags',                   'Array',  'All flags from this run: {type, outcome, severity, detail}'],
    ],
    col_widths=[1.9, 0.8, 3.9],
)

add_heading(doc, 'Collection: health_scores', 2)
add_body(doc, 'One document per component per cycle. Used for trend detection.')
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['_id',         'ObjectId', 'Auto-generated'],
        ['recorded_at', 'String',   'ISO timestamp'],
        ['component',   'String',   '"chatbot", "voice", or "analytics"'],
        ['score',       'Float',    'Severity-weighted health score 0.0–10.0'],
        ['fail_rate',   'Float',    'Fraction of evaluations that failed this cycle'],
        ['details',     'Object',   'chatbot/voice: {total_evaluated, fails}. analytics: {high_flags, medium_flags, total_flags}'],
    ],
    col_widths=[1.4, 0.9, 4.3],
)

add_heading(doc, 'Collection: recommendations', 2)
add_body(doc, 'One document per recommendation cycle.')
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['generated_at',     'String',     'ISO timestamp'],
        ['negative_signals', 'Array[str]', 'Titles of all problems identified this cycle'],
        ['root_cause',       'String',     'Top root cause from highest urgency fix'],
        ['fixes',            'Array',      'Fix objects ranked by urgency: {rank, problem_id, urgency, root_cause, fix, where, change_type, expected_outcome}'],
        ['priority',         'String',     '"HIGH", "MEDIUM", or "LOW" — urgency of top problem'],
    ],
    col_widths=[1.6, 1.1, 3.9],
)

add_heading(doc, 'Collection: feedback', 2)
add_body(doc, 'Human-submitted feedback entries.')
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['text',         'String',  'Raw feedback text'],
        ['submitted_by', 'String',  'Person who submitted'],
        ['submitted_at', 'String',  'ISO timestamp'],
        ['product',      'String',  '"chatbot", "voice_agent", "analytics", or "general"'],
        ['type',         'String',  '"bug", "suggestion", "complaint", or "praise"'],
        ['priority',     'String',  '"high", "medium", or "low"'],
        ['tag',          'String',  'PRICE_ISSUE / RERA_NUMBER / TONE / LANGUAGE / LEAD_CAPTURE / ANALYTICS / META_SIGNAL / PERFORMANCE / OTHER'],
        ['actionable',   'Boolean', 'Whether this feedback can be acted on immediately'],
        ['summary',      'String',  'One-sentence LLM-generated summary'],
    ],
    col_widths=[1.4, 0.9, 4.3],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 14 — LLM WRAPPER
# ════════════════════════════════════════════════════════════════
add_heading(doc, '14.  The LLM Wrapper', 1)
add_body(doc, 'File: llm.py', bold=True)
add_body(doc,
    'All LLM calls go through this single module. It handles the HTTP request to Ollama, '
    'JSON extraction from the response, and retry logic. Every agent calls ask_json() '
    'and receives a Python dict — no agent handles raw LLM output directly.')

add_code(doc,
'ask(prompt: str) → str\n'
'  POST to http://localhost:11434/api/generate\n'
'  Params: model=llama3.1:8b, stream=false, temperature=0.1, num_predict=1500\n'
'  Returns raw text response\n'
'\n'
'_extract_json(text: str) → dict\n'
'  1. Strip markdown code fences (```json ... ```)\n'
'  2. Brace-matching algorithm to find outermost complete { ... } block\n'
'  3. json.loads() and return\n'
'  Raises ValueError if no valid JSON found\n'
'\n'
'ask_json(prompt: str, retries: int = 2) → dict\n'
'  Calls ask(), passes result to _extract_json()\n'
'  Retries up to 2 times if JSON extraction fails\n'
'  Raises on final failure — caught by run_on_new_conversation() which sends error notification'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 15 — NOTIFICATIONS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '15.  Notifications — Teams Webhook', 1)
add_body(doc, 'File: notifier.py', bold=True)

add_heading(doc, 'How It Is Designed to Work', 2)
add_body(doc,
    'When a chatbot conversation or Anandita call fails evaluation, the system fires a '
    'Teams notification immediately. The notification is a structured message containing '
    'the conversation ID, score, summary, number of HIGH severity issues, and the full '
    'issue detail list.')
add_code(doc,
'notify(agent_name, status, details)\n'
'  Prints to console always\n'
'  If TEAMS_WEBHOOK_URL in .env:\n'
'    POST JSON payload to Power Automate webhook URL\n'
'    Power Automate receives it and posts message to Teams channel'
)

add_heading(doc, 'Current Status — Partially Broken', 2)
add_notice(doc,
    '⚠  The webhook fires correctly. Power Automate returns HTTP 202 (accepted). '
    'But messages are not appearing in the Teams channel. '
    'The Python side is working. The Power Automate flow needs to be debugged.',
    'E8691E')

add_body(doc, 'What is confirmed working:', bold=True)
add_bullet(doc, 'notifier.py correctly POSTs to the webhook URL with the right JSON structure')
add_bullet(doc, 'Power Automate receives the request and returns 202')
add_bullet(doc, 'No Python errors — the call succeeds from the code perspective')

add_body(doc, 'What needs to be fixed:', bold=True)
add_bullet(doc, 'Open the Power Automate flow that receives the HTTP trigger')
add_bullet(doc, 'Verify there is a "Post message in a chat or channel" action step after the trigger')
add_bullet(doc, 'Confirm it targets the correct ASBL QA Teams channel')
add_bullet(doc, 'Check that the message body correctly maps to the JSON payload fields sent by notifier.py')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 16 — OUTCOMES CONFIG
# ════════════════════════════════════════════════════════════════
add_heading(doc, '16.  Business Outcomes Configuration', 1)
add_body(doc, 'File: config/outcomes.json', bold=True)
add_body(doc,
    'Business goals are defined in plain English in a JSON config file. '
    'The analytics agent reads this file on every run. No code changes are needed '
    'to add, change, or pause a goal.')

add_heading(doc, 'Currently Active Outcomes', 2)
add_table(doc,
    ['Outcome', 'What It Measures', 'MongoDB Field', 'Overall Target', 'Band5_Hot Target'],
    [
        ['site_visit_booked',      'Lead booked a physical site visit',          'milestones.has_visit_booked',     '30%', '60%'],
        ['otp_verified',           'Lead completed OTP phone verification',      'milestones.has_otp_verified',      '15%', '60%'],
        ['affordability_confirmed','Affordability calculator said YES',           'milestones.affordability_outcome', '10%', '50%'],
    ],
    col_widths=[1.7, 2.0, 2.0, 1.0, 1.0],
)

add_body(doc, 'To change targets or add outcomes — edit config/outcomes.json. Takes effect next run.', italic=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 17 — ASBL PROJECT FACTS
# ════════════════════════════════════════════════════════════════
add_heading(doc, '17.  ASBL Project Facts — Ground Truth', 1)
add_body(doc,
    'These are the authoritative values. If the knowledge base disagrees with this table, '
    'update the KB. Both knowledge_base/webbot/02_kb_project_loft.md AND '
    'knowledge_base/anandita/project_facts.md must be kept in sync.')

add_table(doc,
    ['Project', 'Location', 'Config', 'Sizes & Prices', 'Handover', 'RERA'],
    [
        ['ASBL Loft',
         'Financial District',
         '3BHK only',
         '1695 sqft = ₹1.94 cr + 5% GST\n1870 sqft = ₹2.15 cr + 5% GST\nBHFL booking: ₹10 lakhs\nOther banks: ₹19.4 lakhs\nRental: ₹50/sqft/month till Dec 2026\nNo model flat on site',
         'Dec 2026',
         'P02400006761'],
        ['ASBL Spectra',
         'Financial District',
         '3BHK only',
         '1980 sqft = ₹1.95 cr + 5% GST\n2220 sqft = ₹2.15 cr + 5% GST\nModel flat available',
         'Ready now',
         'P02400003071'],
        ['ASBL Broadway',
         'Financial District',
         '3BHK only',
         '₹9,899 per sqft ONLY\n(Never quote total)',
         'Dec 2029',
         'P02400009684'],
        ['ASBL Landmark',
         'Kukatpally',
         '3BHK & 3.5BHK\n4BHK SOLD OUT',
         '₹8,799 per sqft ONLY\n(Never quote total)',
         'Dec 2028',
         'P02200008770'],
    ],
    col_widths=[1.1, 1.2, 1.0, 2.3, 0.9, 1.1],
)
add_body(doc, 'GST is 5% on all projects. Always disclosed upfront — never hidden in a total figure.', bold=True)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 18 — FILE STRUCTURE
# ════════════════════════════════════════════════════════════════
add_heading(doc, '18.  File & Folder Structure', 1)
add_code(doc,
'asbl-qa-agent/\n'
'├── .env                            MONGO_URI_PROD, MONGO_URI_RESULTS, TEAMS_WEBHOOK_URL\n'
'├── requirements.txt                pymongo, python-dotenv, requests, python-docx\n'
'├── llm.py                          Ollama wrapper — ask(), ask_json(), _extract_json()\n'
'├── db.py                           Results MongoDB — save_*, already_evaluated()\n'
'├── notifier.py                     Console print + Teams webhook POST\n'
'├── runner.py                       4-hour orchestrator — run_once(), loop\n'
'├── generate_doc.py                 Generates this Word document\n'
'│\n'
'├── config/\n'
'│   └── outcomes.json               Business goals — active outcomes and targets\n'
'│\n'
'├── agents/\n'
'│   ├── chatbot_qa.py               Chatbot evaluator — batch(), score_conversation()\n'
'│   ├── voice_qa.py                 Voice evaluator — batch(), score_call()\n'
'│   ├── analytics.py                4 data checks — no LLM\n'
'│   ├── feedback.py                 Health scores + LLM problem detection + feedback CLI\n'
'│   └── recommendation.py           Fix generation — templates + LLM fallback\n'
'│\n'
'├── agents/prompts/\n'
'│   ├── chatbot_qa.md               Evaluation rules — generic, KB-driven, turn-by-turn\n'
'│   ├── voice_qa.md                 Same structure + DECIMAL_NUMBER, LANGUAGE_HANDLING\n'
'│   ├── feedback_reasoning.md       Manager reasoning — cross-agent pattern detection\n'
'│   └── recommendation.md           Fix generation — file locations, known templates\n'
'│\n'
'└── knowledge_base/\n'
'    ├── webbot/                     8 KB files for chatbot evaluation\n'
'    │   ├── 00_qa_checklist.md\n'
'    │   ├── 01_system_prompt.md     Chatbot persona and behavioural rules\n'
'    │   ├── 02_kb_project_loft.md   Loft: prices, RERA, rental offer, payment terms\n'
'    │   ├── 03_kb_market_intelligence.md\n'
'    │   ├── 04_kb_competitive_landscape.md\n'
'    │   ├── 05_kb_persona_playbook.md\n'
'    │   ├── 06_kb_objection_library.md\n'
'    │   ├── 07_kb_resale_framework.md\n'
'    │   └── 08_deployment_guide.md\n'
'    └── anandita/\n'
'        ├── project_facts.md        Prices, RERA, handover for all 4 projects\n'
'        ├── system_prompt.md        Anandita persona + language rules\n'
'        └── qa_checklist.md         Voice evaluation criteria'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 19 — HOW TO RUN
# ════════════════════════════════════════════════════════════════
add_heading(doc, '19.  How to Run', 1)

add_heading(doc, 'Prerequisites', 2)
add_bullet(doc, 'Ollama must be running (check Mac menu bar). Pull the model once: ollama pull llama3.1:8b')
add_bullet(doc, '.env file must have MONGO_URI_PROD, MONGO_URI_RESULTS, and TEAMS_WEBHOOK_URL')
add_bullet(doc, 'Dependencies installed: pip3 install pymongo python-dotenv requests python-docx --break-system-packages')

add_heading(doc, 'Run Commands', 2)
add_code(doc,
'# Full pipeline — one cycle now (testing)\n'
'cd /Users/aanirudhmehra/Desktop/asbl-qa-agent\n'
'python3 runner.py once\n\n'
'# Full pipeline — every 4 hours indefinitely (production)\n'
'python3 runner.py\n\n'
'# Individual agents\n'
'python3 agents/chatbot_qa.py             evaluates last 4 hours of conversations\n'
'python3 agents/chatbot_qa.py 24          evaluates last 24 hours\n'
'python3 agents/voice_qa.py               evaluates last 4 hours of calls\n'
'python3 agents/analytics.py              runs 4 analytics checks\n'
'python3 agents/feedback.py aggregate     health scores + LLM problem detection\n'
'python3 agents/recommendation.py         generates fixes\n\n'
'# Human feedback\n'
'python3 agents/feedback.py submit "Bot gave wrong Spectra price" "Anirudh"\n'
'python3 agents/feedback.py list 20\n\n'
'# Regenerate this document\n'
'python3 generate_doc.py'
)

add_heading(doc, 'What You Can Change Without Any Code', 2)
add_table(doc,
    ['What to Change', 'File to Edit', 'Takes Effect'],
    [
        ['Add a new business goal',              'config/outcomes.json',               'Next run'],
        ['Change target conversion percentages', 'config/outcomes.json',               'Next run'],
        ['Pause a goal',                         'config/outcomes.json → active: false','Next run'],
        ['Change chatbot evaluation rules',      'agents/prompts/chatbot_qa.md',       'Next run'],
        ['Change voice evaluation rules',        'agents/prompts/voice_qa.md',         'Next run'],
        ['Add cross-agent reasoning pattern',    'agents/prompts/feedback_reasoning.md','Next run'],
        ['Add known fix template',               'agents/prompts/recommendation.md',   'Next run'],
        ['Update project prices or RERA',        'knowledge_base/webbot/02_kb_project_loft.md AND knowledge_base/anandita/project_facts.md', 'Next run'],
    ],
    col_widths=[2.4, 2.8, 1.4],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SECTION 20 — BUILT VS PENDING
# ════════════════════════════════════════════════════════════════
add_heading(doc, '20.  What Is Built vs What Is Pending', 1)

add_heading(doc, 'Fully Built and Running', 2)
add_notice(doc, 'Everything below runs automatically every 4 hours without any human intervention.', '1A7A3C')
add_table(doc,
    ['Item', 'Detail'],
    [
        ['Chatbot QA Agent',           'Evaluates all conversations with turnCount ≥ 2. Session signals from events. artifactLabel in evaluation. 8 issue types. PASS/FAIL per conversation.'],
        ['Voice QA Agent',             'Evaluates all call transcripts. DECIMAL_NUMBER check. Language compliance as separate field. 9 issue types including voice-specific ones.'],
        ['Analytics Validator',        '4 checks: band conversion rates, multiplier effectiveness, Meta signal accuracy (0.38%), funnel gaps. Flags stored per cycle.'],
        ['Feedback Agent',             'Health scores with confidence-weighted penalty formula. LLM detects 9 cross-agent patterns. Trend detection over last 3 cycles. Human feedback CLI.'],
        ['Recommendation Agent',       'Template fixes for known issue types. LLM fixes for novel problems. Historical baseline shows when problems started.'],
        ['Runner (4-hour cycle)',       'Orchestrates all 5 agents in sequence. Problems list passed in memory from Feedback to Recommendation.'],
        ['Deduplication',              'already_evaluated() prevents re-processing. Works on conversation_id and call_sid.'],
        ['Results Database (6 colls)', 'chatbot_qa, voice_qa, analytics_runs, health_scores, recommendations, feedback — all fully defined and writing correctly.'],
        ['Chatbot QA Prompt',          'Fully rewritten: generic, KB-driven, turn-by-turn methodology, 3 pre-mortem fixes applied.'],
        ['Voice QA Prompt',            'Rewritten to match chatbot structure with voice-specific additions.'],
        ['Session signals integration','asbl_loft.events queried by sessionId. Ground truth for lead capture, not text inference.'],
    ],
    col_widths=[2.2, 4.4],
    hdr_color=GREEN_HDR,
)

add_heading(doc, 'Partially Built', 2)
add_notice(doc, 'Wired but not working end-to-end.', 'E8691E')
add_table(doc,
    ['Item', 'What Works', 'What Is Broken', 'How to Fix'],
    [
        ['Teams Webhook',
         'notifier.py POSTs correctly. Power Automate returns 202.',
         'Message does not appear in Teams channel.',
         'Open the Power Automate flow. Verify "Post message in channel" action exists and targets the right channel. Check message body mapping.'],
    ],
    col_widths=[1.4, 1.8, 1.8, 1.6],
    hdr_color=ORANGE_HDR,
)

add_heading(doc, 'Pending — Not Yet Built', 2)
add_notice(doc, 'Designed or planned but not yet implemented.', 'C02020')
add_table(doc,
    ['#', 'Item', 'Priority', 'What Is Needed', 'Dependency'],
    [
        ['1', 'End-to-end test run',
         'HIGH',
         'Run python3 runner.py once against real data. Validate that updated chatbot and voice prompts produce correct results on real conversations.',
         'None — can run today'],
        ['2', 'Voice QA context enrichment',
         'MEDIUM',
         'Pass call_outcome, site_visit_agreed, ready_to_book as structured context above the transcript in voice_qa.py. These fields exist in call_transcripts and give the LLM richer signal.',
         'Code change in voice_qa.py only'],
        ['3', 'NO_ANSWER call explicit filter',
         'LOW',
         'Add call_outcome == "NO_ANSWER" filter in voice_qa.batch() instead of relying on empty transcript check.',
         'Small code change in voice_qa.py'],
        ['4', 'Analytics site_visits cross-check',
         'MEDIUM',
         'Cross-reference analytics_db.site_visits confirmed visit records against milestones.has_visit_booked in scores_overall. Flag any mismatches.',
         'Code change in analytics.py'],
        ['5', 'Security / code check',
         'MEDIUM',
         'Add bitbucket-pipelines.yml to website and voice agent Bitbucket repos. Runs Bandit (security), pip-audit (dependencies), pylint (quality) on every PR. Completely free using Bitbucket Pipelines.',
         'Write access to both Bitbucket repos'],
        ['6', 'Security Agent (QA integration)',
         'LOW',
         'agents/security.py reads scan results and surfaces HIGH findings through Teams and qa_results database.',
         'Bitbucket App Password OR CodeAnt API key'],
        ['7', 'Meta Agent',
         'LOW',
         'Dedicated agent for deeper Meta ad performance analysis beyond the current accuracy check.',
         'Deliberately deferred to next cycle'],
    ],
    col_widths=[0.3, 1.5, 0.8, 3.0, 1.0],
    hdr_color=RED_HDR,
)

# ── Save ──────────────────────────────────────────────────────────
output_path = '/Users/aanirudhmehra/Desktop/asbl_qa_elaborated.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
